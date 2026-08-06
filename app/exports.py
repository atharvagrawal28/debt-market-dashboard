"""Report renderers: TXT, HTML, CSV, DOCX, XLSX.

Every renderer walks the same section list, so a new field added to the report
schema shows up in all five formats without touching each one.
"""
import csv
import html as html_mod
import io

# Footer line carried on every rendered report, matching the dashboards.
CREDIT = "Made By Finqrate for Swayamjit Dalai"
FINQRATE_URL = "https://www.linkedin.com/company/finqrate"
SWAYAMJIT_URL = "https://www.linkedin.com/in/swayamjitdalai/"
CREDIT_LINKS = f"Finqrate: {FINQRATE_URL}  |  Swayamjit Dalai: {SWAYAMJIT_URL}"

# (heading, key, kind, columns) -- columns are (header, field, format)
NUM = "{:.4f}"
PCT = "{:.2f}"

SECTIONS = [
    ("Money Market", "mm_detail", "table", [
        ("Instrument", "kind", None), ("Open", "open", PCT), ("High", "high", PCT),
        ("Low", "low", PCT), ("Last Trade", "last_trade", PCT),
        ("Wtd Avg", "weighted_avg", NUM), ("Volume (Cr)", "volume_cr", "{:,.2f}")]),
    ("Benchmark G-Sec Curve", "benchmark_curve", "table", [
        ("Tenor", "tenor_label", None), ("Security", "security_name", None),
        ("Yield", "yield_today", NUM), ("Prev", "yield_prev", NUM),
        ("Chg (bps)", "change_bps", "{:+.1f}")]),
    ("Top Traded G-Secs", "top_traded_gsecs", "table", [
        ("Security", "name", None), ("Yield", "yield_today", NUM),
        ("Prev", "yield_prev", NUM)]),
    ("T-Bill", "tbill_range", "table", [
        ("Tenor", "tenor_label", None), ("Security", "security_name", None),
        ("Rate", "rate_today", NUM), ("Prev", "rate_prev", NUM),
        ("Chg (bps)", "change_bps", "{:+.1f}")]),
    ("SDL", "sdl_range", "table", [
        ("Tenor", "tenor_label", None), ("Security", "security_name", None),
        ("Rate", "rate_today", NUM), ("Prev", "rate_prev", NUM)]),
    ("CD Secondary - Tenor Ranges", "cd_money_market", "table", [
        ("Tenor", "tenor_label", None), ("Low", "low", PCT), ("High", "high", PCT),
        ("Wtd Avg", "weighted_avg", NUM), ("Volume (Cr)", "volume_cr", "{:,.2f}"),
        ("Trades", "trade_count", None), ("Issuers", "issuer_count", None)]),
    ("CP Secondary - Tenor Ranges", "cp_money_market", "table", [
        ("Tenor", "tenor_label", None), ("Low", "low", PCT), ("High", "high", PCT),
        ("Wtd Avg", "weighted_avg", NUM), ("Volume (Cr)", "volume_cr", "{:,.2f}"),
        ("Trades", "trade_count", None), ("Issuers", "issuer_count", None)]),
    ("Spread over T-Bill", "spreads", "table", [
        ("Tenor", "tenor_label", None), ("T-Bill", "tbill_ytm", NUM),
        ("CD Wtd Avg", "cd_wavg", NUM), ("CD Spread (bps)", "cd_spread_bps", "{:+.1f}"),
        ("CP Wtd Avg", "cp_wavg", NUM), ("CP Spread (bps)", "cp_spread_bps", "{:+.1f}")]),
    ("AAA PSU Corporate Bonds", "aaa_psu_corp", "table", [
        ("Tenor", "tenor_label", None), ("Low", "low", PCT), ("High", "high", PCT)]),
    ("OIS Curve", "ois_curve", "table", [
        ("Tenor", "tenor_label", None), ("Rate", "rate_today", NUM),
        ("Prev", "rate_prev", NUM)]),
    ("Most Active CD Issuers", "cd_issuers", "table", [
        ("Issuer", "issuer", None), ("Volume (Cr)", "volume_cr", "{:,.2f}"),
        ("Low", "low", PCT), ("High", "high", PCT), ("Wtd Avg", "weighted_avg", NUM)]),
    ("Most Active CP Issuers", "cp_issuers", "table", [
        ("Issuer", "issuer", None), ("Volume (Cr)", "volume_cr", "{:,.2f}"),
        ("Low", "low", PCT), ("High", "high", PCT), ("Wtd Avg", "weighted_avg", NUM)]),
    ("News", "news", "table", [
        ("Category", "category", None), ("Headline", "headline", None)]),
]


def _fmt(value, spec):
    if value is None or value == "":
        return "-"
    if spec:
        try:
            return spec.format(value)
        except (ValueError, TypeError):
            return str(value)
    return str(value)


def _rows(report, key, columns):
    data = report.get(key) or []
    if isinstance(data, dict):
        data = [data]
    return [[_fmt(item.get(field), spec) for _, field, spec in columns]
            for item in data if isinstance(item, dict)]


def _scalars(report):
    fx = report.get("fx") or {}
    brent = report.get("brent") or {}
    out = []
    if fx.get("close"):
        out.append(("USD/INR Close", _fmt(fx.get("close"), "{:.4f}")))
        if fx.get("open"):
            out.append(("USD/INR Open", _fmt(fx.get("open"), "{:.4f}")))
        if fx.get("day_low") or fx.get("day_high"):
            out.append(("USD/INR Range",
                        f"{_fmt(fx.get('day_low'), NUM)} - {_fmt(fx.get('day_high'), NUM)}"))
    if brent.get("price_usd"):
        out.append(("Brent (USD)", _fmt(brent.get("price_usd"), "{:.2f}")))
    return out


def to_txt(report):
    date = report.get("report_date", "")
    buf = [f"FIXED INCOME CLOSING REPORT - {date}", "=" * 64, ""]
    for label, value in _scalars(report):
        buf.append(f"{label:<22} {value}")
    if _scalars(report):
        buf.append("")
    for heading, key, _kind, columns in SECTIONS:
        rows = _rows(report, key, columns)
        if not rows:
            continue
        headers = [h for h, _, _ in columns]
        widths = [max(len(headers[i]), *(len(r[i]) for r in rows))
                  for i in range(len(headers))]
        buf.append(heading.upper())
        buf.append("-" * (sum(widths) + 3 * len(widths)))
        buf.append("   ".join(h.ljust(widths[i]) for i, h in enumerate(headers)))
        for r in rows:
            buf.append("   ".join(c.ljust(widths[i]) for i, c in enumerate(r)))
        buf.append("")
    commentary = (report.get("ai_commentary") or "").strip()
    if commentary:
        buf += ["CLOSING COMMENTARY", "-" * 64, commentary, ""]
    sources = report.get("_sources") or []
    if sources:
        buf.append("Sources: " + "; ".join(dict.fromkeys(sources)))
    buf += ["", "-" * 64, CREDIT, CREDIT_LINKS]
    return "\n".join(buf)


def to_html(report):
    e = html_mod.escape
    date = report.get("report_date", "")
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>Fixed Income Closing Report {e(date)}</title>",
        "<style>body{font-family:Segoe UI,Arial,sans-serif;margin:32px;color:#0f172a}"
        "h1{font-size:22px;margin:0 0 4px}h2{font-size:14px;text-transform:uppercase;"
        "letter-spacing:.04em;color:#1d4ed8;margin:28px 0 8px}"
        "table{border-collapse:collapse;font-size:13px;width:100%;max-width:900px}"
        "th{background:#1d4ed8;color:#fff;padding:7px 9px;text-align:left;font-size:11px;"
        "text-transform:uppercase;letter-spacing:.03em}"
        "td{padding:6px 9px;border-bottom:1px solid #e2e8f0}"
        "tr:nth-child(even) td{background:#f8fafc}"
        ".kv{font-size:13px;color:#334155}.src{margin-top:28px;font-size:12px;color:#64748b}"
        "</style></head><body>",
        f"<h1>Fixed Income Closing Report</h1><div class='kv'>{e(date)}</div>",
    ]
    for label, value in _scalars(report):
        parts.append(f"<div class='kv'><b>{e(label)}:</b> {e(value)}</div>")
    for heading, key, _kind, columns in SECTIONS:
        rows = _rows(report, key, columns)
        if not rows:
            continue
        parts.append(f"<h2>{e(heading)}</h2><table><tr>")
        parts += [f"<th>{e(h)}</th>" for h, _, _ in columns]
        parts.append("</tr>")
        for r in rows:
            parts.append("<tr>" + "".join(f"<td>{e(c)}</td>" for c in r) + "</tr>")
        parts.append("</table>")
    commentary = (report.get("ai_commentary") or "").strip()
    if commentary:
        parts.append("<h2>Closing Commentary</h2><p>" +
                     e(commentary).replace("\n", "<br>") + "</p>")
    sources = list(dict.fromkeys(report.get("_sources") or []))
    if sources:
        parts.append("<div class='src'>Sources: " + e("; ".join(sources)) + "</div>")
    parts.append(
        "<div style='margin-top:26px;padding-top:14px;border-top:1px solid #e2e8f0;"
        "text-align:center;font-size:12px;color:#64748b;letter-spacing:.04em'>"
        f"Made By <a href='{FINQRATE_URL}' style='color:#1d4ed8;font-weight:700;"
        "text-decoration:none'>Finqrate</a> for "
        f"<a href='{SWAYAMJIT_URL}' style='color:#334155;text-decoration:none'>"
        "Swayamjit Dalai</a></div>")
    parts.append("</body></html>")
    return "".join(parts)


def to_csv(report):
    out = io.StringIO()
    w = csv.writer(out, lineterminator="\n")
    w.writerow(["Fixed Income Closing Report", report.get("report_date", "")])
    w.writerow([])
    for label, value in _scalars(report):
        w.writerow([label, value])
    for heading, key, _kind, columns in SECTIONS:
        rows = _rows(report, key, columns)
        if not rows:
            continue
        w.writerow([])
        w.writerow([heading])
        w.writerow([h for h, _, _ in columns])
        w.writerows(rows)
    commentary = (report.get("ai_commentary") or "").strip()
    if commentary:
        w.writerow([])
        w.writerow(["Closing Commentary"])
        w.writerow([commentary])
    w.writerow([])
    w.writerow([CREDIT])
    w.writerow([FINQRATE_URL, SWAYAMJIT_URL])
    return out.getvalue()


def to_docx(report):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    heading = doc.add_heading("Fixed Income Closing Report", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    doc.add_paragraph(report.get("report_date", ""))

    scalars = _scalars(report)
    if scalars:
        for label, value in scalars:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)

    for section_heading, key, _kind, columns in SECTIONS:
        rows = _rows(report, key, columns)
        if not rows:
            continue
        doc.add_heading(section_heading, level=2)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Light Grid Accent 1"
        for i, (label, _, _) in enumerate(columns):
            cell = table.rows[0].cells[i]
            cell.text = label
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for r in rows:
            cells = table.add_row().cells
            for i, value in enumerate(r):
                cells[i].text = value
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9)

    commentary = (report.get("ai_commentary") or "").strip()
    if commentary:
        doc.add_heading("Closing Commentary", level=2)
        doc.add_paragraph(commentary)

    sources = list(dict.fromkeys(report.get("_sources") or []))
    if sources:
        p = doc.add_paragraph()
        run = p.add_run("Sources: " + "; ".join(sources))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CREDIT)
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CREDIT_LINKS)
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_xlsx(report):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Closing Report"

    header_fill = PatternFill("solid", fgColor="1D4ED8")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    title_font = Font(bold=True, size=14, color="1D4ED8")
    section_font = Font(bold=True, size=11, color="1D4ED8")
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.cell(row=1, column=1, value="Fixed Income Closing Report").font = title_font
    ws.cell(row=2, column=1, value=report.get("report_date", ""))
    row = 4

    for label, value in _scalars(report):
        ws.cell(row=row, column=1, value=label).font = Font(bold=True, size=10)
        ws.cell(row=row, column=2, value=value)
        row += 1
    if _scalars(report):
        row += 1

    widths = {}
    for heading, key, _kind, columns in SECTIONS:
        data = report.get(key) or []
        if isinstance(data, dict):
            data = [data]
        data = [d for d in data if isinstance(d, dict)]
        if not data:
            continue
        ws.cell(row=row, column=1, value=heading).font = section_font
        row += 1
        for i, (label, _, _) in enumerate(columns, start=1):
            c = ws.cell(row=row, column=i, value=label)
            c.fill, c.font, c.border = header_fill, header_font, border
            c.alignment = Alignment(horizontal="center")
            widths[i] = max(widths.get(i, 10), len(label) + 2)
        row += 1
        for item in data:
            for i, (_, field, spec) in enumerate(columns, start=1):
                value = item.get(field)
                cell = ws.cell(row=row, column=i, value=value)
                cell.border = border
                if isinstance(value, (int, float)) and spec:
                    cell.number_format = ("#,##0.00" if "," in spec
                                          else "0.0000" if "4f" in spec else "0.00")
                    cell.alignment = Alignment(horizontal="right")
                widths[i] = max(widths.get(i, 10), len(str(value if value is not None else "")) + 2)
            row += 1
        row += 1

    commentary = (report.get("ai_commentary") or "").strip()
    if commentary:
        ws.cell(row=row, column=1, value="Closing Commentary").font = section_font
        cell = ws.cell(row=row + 1, column=1, value=commentary)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        row += 3

    credit = ws.cell(row=row + 1, column=1, value=CREDIT)
    credit.font = Font(bold=True, size=9, color="1D4ED8")
    links = ws.cell(row=row + 2, column=1, value=CREDIT_LINKS)
    links.font = Font(size=8, color="64748B")

    for i, width in widths.items():
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = min(width, 42)
    ws.freeze_panes = "A4"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


RENDERERS = {
    "txt": (to_txt, "text/plain; charset=utf-8", "txt"),
    "html": (to_html, "text/html; charset=utf-8", "html"),
    "csv": (to_csv, "text/csv; charset=utf-8", "csv"),
    "docx": (to_docx,
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
             "docx"),
    "xlsx": (to_xlsx,
             "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
             "xlsx"),
}


def render(report, fmt):
    if fmt not in RENDERERS:
        raise ValueError(f"unknown format: {fmt}")
    fn, mimetype, ext = RENDERERS[fmt]
    payload = fn(report)
    if isinstance(payload, str):
        payload = payload.encode("utf-8")
    return payload, mimetype, ext


def trades_xlsx(rows, title="Trades"):
    """The CD/CP dashboard's download, rendered server-side with real styling.

    SheetJS's community build silently drops cell styles, so the browser-side
    download came out unformatted despite setting them.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]

    headers = ["DEAL DATE", "ISSUER", "INSTRUMENT", "MATURITY", "RATE (%)", "QTM (Crs)"]
    fill = PatternFill("solid", fgColor="FFFF00")
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for i, h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=i, value=h)
        c.font = Font(bold=True, size=11)
        c.fill = fill
        c.border = border
        c.alignment = Alignment(horizontal="center", vertical="center")

    for r, row in enumerate(rows, start=2):
        values = [row.get("deal_date"), row.get("issuer"), row.get("instrument"),
                  row.get("maturity_date"), row.get("yield_pct"), row.get("amount_cr")]
        for i, value in enumerate(values, start=1):
            c = ws.cell(row=r, column=i, value=value)
            c.border = border
            c.alignment = Alignment(horizontal="center", vertical="center")
            if i == 5:
                c.number_format = "0.00"
            if i == 6:
                c.number_format = "#,##0.00"

    credit = ws.cell(row=len(rows) + 3, column=1, value=CREDIT)
    credit.font = Font(bold=True, size=9, color="1D4ED8")
    links = ws.cell(row=len(rows) + 4, column=1, value=CREDIT_LINKS)
    links.font = Font(size=8, color="64748B")

    for i, width in enumerate([14, 34, 12, 14, 11, 13], start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = width
    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
